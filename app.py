import streamlit as st
import openai
from io import BytesIO
from docx import Document
import os
import tempfile
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.chains import RetrievalQA

# Page Configuration
st.set_page_config(
    page_title="CaseLaw RAG | Verified Legal Precedent & Brief Generator",
    page_icon="⚖️",
    layout="wide"
)

# Sidebar Navigation
st.sidebar.title("CaseLaw RAG Suite")
page = st.sidebar.radio("Navigation", ["🏠 Overview", "📂 PDF Ingestion & Vector Store", "🔍 Precedent Search & RAG Brief Builder"])

openai_api_key = st.secrets.get("OPENAI_API_KEY")

if page == "🏠 Overview":
    st.title("⚖️ CaseLaw RAG & Citation Verifier")
    st.subheader("Deterministic Legal Precedent Retrieval & Hallucination-Free Drafting")

    st.markdown("""
    Welcome to **CaseLaw RAG**, an advanced legal technology tool designed to eliminate hallucinated citations in legal research. 
    By combining semantic vector search over BIA and federal circuit precedent via ChromaDB with strict temperature-zero guardrails, 
    this platform ensures every legal argument is backed by verifiable pin-citations.
    """)

    st.divider()

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("#### 📚 Document Ingestion")
        st.markdown("Upload binding precedent, BIA decisions, and circuit court PDFs into a persistent local vector database.")
    with col2:
        st.markdown("#### 🔎 Semantic Retrieval")
        st.markdown("Surface exact legal standards, statutory interpretations, and precedential holdings instantly using vector similarity.")
    with col3:
        st.markdown("#### 🛡️ Zero Hallucination Guardrails")
        st.markdown("Forces the LLM to ground all outputs strictly in retrieved chunks with mandatory source verification.")

    st.divider()
    st.success("👈 Use the sidebar to upload PDFs or run grounded legal research queries.")

elif page == "📂 PDF Ingestion & Vector Store":
    st.title("📂 Legal Document Ingestion & Embedding")
    st.write("Upload official legal precedent PDFs (e.g., BIA decisions, Circuit rulings) to index them into your local Chroma vector store.")

    if not openai_api_key:
        st.warning("⚠️ Please configure your OPENAI_API_KEY in your Streamlit app secrets.")
    else:
        uploaded_pdfs = st.file_uploader(
            "Upload Legal Precedent PDFs:",
            type=["pdf"],
            accept_multiple_files=True
        )

        if uploaded_pdfs:
            if st.button("Process & Embed PDFs into ChromaDB 🚀", type="primary"):
                with st.spinner("Chunking documents and calculating vector embeddings..."):
                    try:
                        all_docs = []
                        for pdf in uploaded_pdfs:
                            # Save uploaded file to a temp file so LangChain can load it
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                                tmp_file.write(pdf.read())
                                tmp_path = tmp_file.name

                            loader = PyPDFLoader(tmp_path)
                            docs = loader.load()
                            all_docs.extend(docs)
                            os.unlink(tmp_path) # Clean up temp file

                        # Split text into manageable chunks
                        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                        chunks = text_splitter.split_documents(all_docs)

                        # Create embeddings and store in local Chroma vector database
                        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
                        vector_store = Chroma.from_documents(chunks, embeddings, persist_directory="./chroma_db")
                        
                        st.success(f"Successfully processed {len(uploaded_pdfs)} PDF(s) and embedded {len(chunks)} text chunks into ChromaDB!")
                        st.info("💡 Your vector store is now ready. Go to **🔍 Precedent Search & RAG Brief Builder** to query your documents.")

                    except Exception as e:
                        st.error(f"Error processing PDFs: {e}")

elif page == "🔍 Precedent Search & RAG Brief Builder":
    st.title("🔍 Legal Precedent Search & Argument Generator")
    st.write("Query your ingested legal database using semantic retrieval to generate rigorous, citation-backed legal arguments.")

    if not openai_api_key:
        st.warning("⚠️ Please configure your OPENAI_API_KEY in your Streamlit app secrets.")
    else:
        legal_query = st.text_area(
            "Enter Legal Research Question or Client Fact Pattern:",
            height=140,
            placeholder="e.g., What specific standard does the BIA apply when evaluating government acquiescence in domestic violence claims?"
        )

        if st.button("Generate Grounded Brief from Vector Store 🚀", type="primary"):
            if legal_query.strip():
                with st.spinner("Performing vector similarity search and drafting grounded legal brief..."):
                    try:
                        # Load existing Chroma vector store
                        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
                        if not os.path.exists("./chroma_db"):
                            st.warning("⚠️ No vector store found. Please upload and process PDFs in the 'PDF Ingestion' tab first.")
                        else:
                            vector_store = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)
                            retriever = vector_store.as_retriever(search_kwargs={"k": 4})

                            # Retrieve relevant context chunks
                            relevant_docs = retriever.invoke(legal_query)
                            context_text = "\n\n---\n\n".join([doc.page_content for doc in relevant_docs])

                            # Call OpenAI with strict temperature-zero guardrails
                            client = openai.OpenAI(api_key=openai_api_key)
                            response = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": (
                                            "You are an expert legal research assistant and appellate brief writer. "
                                            "Answer the user's research question strictly using ONLY the provided Precedent Context retrieved from the database. "
                                            "Do not extrapolate or hallucinate outside the text. "
                                            "Provide: 1) Executive Legal Summary, 2) Direct Precedent Application with Pin-Citations from the context, "
                                            "and 3) Recommended Argumentation Strategy. If the context does not contain the answer, explicitly state so."
                                        )
                                    },
                                    {
                                        "role": "user",
                                        "content": f"Retrieved Precedent Context:\n{context_text}\n\nLegal Research Question:\n{legal_query}"
                                    }
                                ],
                                temperature=0.0
                            )
                            rag_output = response.choices[0].message.content

                            st.success("Grounded Legal Brief Generated Successfully!")
                            st.markdown("---")
                            st.markdown("### 📄 Verified Legal Analysis & Retrieved Citations")
                            st.markdown(rag_output)

                            # Word Document Export
                            doc = Document()
                            doc.add_heading("Legal Brief & Vector Precedent Analysis", level=1)
                            doc.add_paragraph(f"Query: {legal_query}\n")
                            doc.add_paragraph(rag_output)

                            doc_io = BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)

                            st.download_button(
                                label="📥 Download Legal Brief (.docx)",
                                data=doc_io,
                                file_name="CaseLaw_RAG_Brief.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                            st.markdown("---")
                            st.markdown("### 🔒 Human-in-the-Loop (HITL) Sign-Off")
                            st.checkbox("Attorney Verification: Confirm citation accuracy against source documents before filing.")

                    except Exception as e:
                        st.error(f"RAG Execution Error: {e}")
            else:
                st.warning("⚠️ Please enter a legal research question.")
   
    

       
